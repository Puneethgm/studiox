import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

# Create test directory
dir_path = "/home/infaira/Desktop/studiox/test_documents"
os.makedirs(dir_path, exist_ok=True)

# 1. TEXT FAQ File
faq_text = """YOGA STUDIO FREQUENTLY ASKED QUESTIONS (FAQ)

Q: What should I wear to a class?
A: Wear comfortable, stretchy clothing. We recommend bare feet for all yoga styles.

Q: Do you provide mats?
A: Yes! We provide complimentary mats, blocks, and straps. You are also welcome to bring your own.

Q: Can beginners join?
A: Absolutely! We welcome all levels. "Yoga Basics" and "Gentle Flow" are excellent classes to start with.

Q: What is your cancellation policy?
A: Please cancel at least 4 hours before class starts to avoid a late charge.
"""

with open(os.path.join(dir_path, "yoga_faq.txt"), "w") as f:
    f.write(faq_text)

# 2. CSV Schedule File
csv_data = """Day,Class Name,Time,Instructor,Difficulty
Monday,Vinyasa Flow,07:00 AM - 08:00 AM,Sarah Jenkins,Intermediate
Monday,Yoga Basics,09:30 AM - 10:30 AM,Michael Chang,Beginner
Tuesday,Yin Yoga,06:00 PM - 07:15 PM,Elena Rostova,All Levels
Wednesday,Power Yoga,07:00 AM - 08:00 AM,Sarah Jenkins,Advanced
Thursday,Gentle Flow,10:00 AM - 11:00 AM,Michael Chang,Beginner
Friday,Hot Yoga,05:30 PM - 06:45 PM,Elena Rostova,Intermediate
"""

with open(os.path.join(dir_path, "yoga_schedule.csv"), "w") as f:
    f.write(csv_data)

# 3. DOCX Pricing File
doc = Document()
doc.add_heading("YOGA STUDIO PRICING AND PLANS", level=0)

doc.add_heading("Introductory Offer", level=1)
doc.add_paragraph("New members enjoy our Intro Pass: 3 Classes for $30 (expires 14 days from purchase).")

doc.add_heading("Memberships", level=1)
doc.add_paragraph("• Monthly Unlimited: $120/month (autopay, 3-month minimum commitment). Includes 10% off workshops.")
doc.add_paragraph("• Annual Pass: $1,200/year (save $240).")

doc.add_heading("Class Passes", level=1)
doc.add_paragraph("• Single Class Drop-in: $22")
doc.add_paragraph("• 10-Class Card: $180 (expires in 6 months)")

doc.save(os.path.join(dir_path, "yoga_pricing.docx"))

# 4. PDF Rules File
pdf_path = os.path.join(dir_path, "yoga_rules.pdf")
pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
styles = getSampleStyleSheet()

story = []
story.append(Paragraph("<b>YOGA STUDIO POLICIES AND HOUSE RULES</b>", styles["Title"]))
story.append(Spacer(1, 20))

rules = [
    "<b>1. Arrive Early:</b> Please arrive 10-15 minutes before class starts. Late entry is not permitted to preserve the peaceful environment.",
    "<b>2. Silence Your Devices:</b> Turn off or silence all mobile phones and smartwatches before entering the studio space.",
    "<b>3. Respect the Space:</b> Leave your shoes and belongings in the cubbies outside the main studio room.",
    "<b>4. Clean Your Gear:</b> If using studio mats or blocks, please wipe them down with the provided organic disinfectant spray after class."
]

for rule in rules:
    story.append(Paragraph(rule, styles["Normal"]))
    story.append(Spacer(1, 10))

pdf_doc.build(story)

print("Generated all files successfully in:", dir_path)
