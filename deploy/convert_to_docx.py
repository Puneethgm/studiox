import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

def md_to_docx():
    md_path = 'deploy/complete_ec2_split_deployment_guide.md'
    docx_path = 'deploy/complete_ec2_split_deployment_guide.docx'
    
    doc = Document()
    
    # Custom styles / colors
    COLOR_PRIMARY = RGBColor(31, 78, 121)   # Blue
    COLOR_CODE = RGBColor(199, 37, 78)     # Pinkish red for inline code
    COLOR_DARK = RGBColor(51, 51, 51)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = COLOR_DARK
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.split('\n')
    
    in_code_block = False
    code_text = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Code block handler
        if line.strip().startswith('```'):
            if in_code_block:
                # Close code block
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.space_before = Pt(6)
                
                # Add shading to the paragraph
                pPr = p._p.get_or_add_pPr()
                shading = parse_xml(r'<w:shd {} w:fill="F4F4F6"/>'.format(nsdecls('w')))
                pPr.append(shading)
                
                run = p.add_run('\n'.join(code_text))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(40, 40, 40)
                code_text = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_text.append(line)
            i += 1
            continue
            
        # Headers
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            h.runs[0].font.color.rgb = COLOR_PRIMARY
            i += 1
            continue
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            h.runs[0].font.color.rgb = COLOR_PRIMARY
            i += 1
            continue
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
            h.runs[0].font.color.rgb = COLOR_PRIMARY
            i += 1
            continue
        elif line.startswith('#### '):
            h = doc.add_heading(line[5:], level=4)
            h.runs[0].font.color.rgb = COLOR_PRIMARY
            i += 1
            continue
            
        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(r'<w:pBdr {}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>'.format(nsdecls('w')))
            pPr.append(pBdr)
            i += 1
            continue
            
        # Table handler
        if line.strip().startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|') and '-' in lines[i+1]:
            # Read table
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            # Parse table rows (skip separator row at index 1)
            rows_data = []
            for idx, tl in enumerate(table_lines):
                if idx == 1:
                    continue
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                rows_data.append(cells)
                
            if rows_data:
                num_cols = len(rows_data[0])
                table = doc.add_table(rows=0, cols=num_cols)
                table.style = 'Light Shading Accent 1'
                for row_idx, rdata in enumerate(rows_data):
                    row_cells = table.add_row().cells
                    for col_idx, text in enumerate(rdata):
                        if col_idx < len(row_cells):
                            # Remove bold/code markdown tags in cells
                            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                            clean_text = re.sub(r'`(.*?)`', r'\1', clean_text)
                            row_cells[col_idx].text = clean_text
            continue
            
        # Blockquote / Alert
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            pPr = p._p.get_or_add_pPr()
            # Add left border shading
            pBdr = parse_xml(r'<w:pBdr {}><w:left w:val="single" w:sz="24" w:space="8" w:color="007ACC"/></w:pBdr>'.format(nsdecls('w')))
            pPr.append(pBdr)
            shading = parse_xml(r'<w:shd {} w:fill="F0F7FF"/>'.format(nsdecls('w')))
            pPr.append(shading)
            
            cleaned = line[2:]
            if cleaned.startswith('[!'):
                # Strip out alert tags like [!IMPORTANT]
                cleaned = re.sub(r'^\[\![A-Z]+\]\s*', '', cleaned)
            
            run = p.add_run(cleaned)
            run.font.italic = True
            run.font.size = Pt(10)
            i += 1
            continue
            
        # Lists
        is_bullet = line.strip().startswith('* ') or line.strip().startswith('- ')
        is_num = re.match(r'^\s*\d+\.\s', line) is not None
        
        if is_bullet or is_num:
            style_name = 'List Bullet' if is_bullet else 'List Number'
            cleaned = re.sub(r'^(\s*\*|\s*-|\s*\d+\.)\s*', '', line)
            p = doc.add_paragraph(style=style_name)
            
            # Format bold, code tags inline
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', cleaned)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.color.rgb = COLOR_CODE
                else:
                    p.add_run(part)
            i += 1
            continue
            
        # Normal paragraphs
        if line.strip() != "":
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.color.rgb = COLOR_CODE
                else:
                    p.add_run(part)
        
        i += 1
        
    doc.save(docx_path)
    print("Successfully converted md to docx!")

if __name__ == '__main__':
    md_to_docx()
